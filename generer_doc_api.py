"""
Script de génération du document Word expliquant les APIs
Google Identity Services et Stripe utilisées dans le projet Tennis Club.

Usage :
    cd d:\ProjetIntegrationDev\TennisClub
    .\env\Scripts\Activate.ps1
    python generer_doc_api.py

Génère : exemples/Documentation_APIs.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def ajouter_titre_page(doc):
    """Page de titre"""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documentation des APIs")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xC1, 0x44, 0x0E)  # terre-battue

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Google Identity Services & Stripe")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2D, 0x50, 0x16)  # vert-court

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Projet Tennis Club — Intégration de dispositifs")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fouad El Kadouri")
    run.font.size = Pt(12)

    doc.add_page_break()


def ajouter_table_des_matieres(doc):
    """Table des matières manuelle"""
    doc.add_heading("Table des matières", level=1)

    items = [
        ("1.", "Introduction"),
        ("2.", "API Google Identity Services"),
        ("  2.1.", "Qu'est-ce que c'est ?"),
        ("  2.2.", "Comment obtenir les clés API"),
        ("  2.3.", "Comment ça marche"),
        ("  2.4.", "Intégration dans le Tennis Club"),
        ("  2.5.", "Exemple de code"),
        ("3.", "API Stripe Checkout"),
        ("  3.1.", "Qu'est-ce que c'est ?"),
        ("  3.2.", "Comment obtenir les clés API"),
        ("  3.3.", "Comment ça marche"),
        ("  3.4.", "Intégration dans le Tennis Club"),
        ("  3.5.", "Exemple de code"),
        ("4.", "Résumé comparatif"),
    ]

    for num, titre in items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {titre}")
        run.font.size = Pt(11)
        if not num.startswith(" "):
            run.bold = True

    doc.add_page_break()


def ajouter_introduction(doc):
    """Section 1 : Introduction"""
    doc.add_heading("1. Introduction", level=1)

    doc.add_paragraph(
        "Ce document explique les deux APIs externes utilisées dans le projet "
        "Tennis Club pour la gestion des terrains de tennis :"
    )

    doc.add_paragraph(
        "Google Identity Services — permet aux membres de se connecter avec "
        "leur compte Google pour lier leur email avant de payer.",
        style='List Bullet'
    )

    doc.add_paragraph(
        "Stripe Checkout — permet aux membres de payer leur cotisation annuelle "
        "en ligne de manière sécurisée (carte bancaire).",
        style='List Bullet'
    )

    doc.add_paragraph(
        "Pour chaque API, nous expliquons : à quoi elle sert, comment obtenir les "
        "clés nécessaires, comment elle fonctionne techniquement, comment elle est "
        "intégrée dans notre site, et un exemple de code complet."
    )

    doc.add_page_break()


def ajouter_section_google(doc):
    """Section 2 : Google Identity Services"""
    doc.add_heading("2. API Google Identity Services", level=1)

    # 2.1
    doc.add_heading("2.1. Qu'est-ce que c'est ?", level=2)
    doc.add_paragraph(
        "Google Identity Services (GIS) est un service gratuit de Google qui permet "
        "à n'importe quel site web d'afficher un bouton « Se connecter avec Google ». "
        "Quand l'utilisateur clique dessus, une fenêtre Google s'ouvre, il entre son "
        "email et mot de passe Google, et Google renvoie un « jeton » (token) contenant "
        "les informations de l'utilisateur (email, nom, photo de profil)."
    )
    doc.add_paragraph(
        "C'est le même système utilisé par des sites comme Spotify, Airbnb ou Canva "
        "quand ils proposent « Continuer avec Google »."
    )

    # 2.2
    doc.add_heading("2.2. Comment obtenir les clés API", level=2)
    doc.add_paragraph("Pour utiliser Google Identity Services, il faut :")

    etapes_google = [
        ("Créer un projet Google Cloud",
         "Aller sur https://console.cloud.google.com, créer un nouveau projet "
         "(ex: « Tennis Club »)."),
        ("Activer l'API",
         "Dans le menu « API et services » → « Bibliothèque », chercher "
         "« Google Identity » et l'activer."),
        ("Créer un identifiant OAuth",
         "Dans « API et services » → « Identifiants » → « Créer des identifiants » "
         "→ « ID client OAuth 2.0 ». Choisir « Application Web » et ajouter "
         "http://localhost:8000 dans les origines autorisées."),
        ("Récupérer le Client ID",
         "Google affiche un identifiant qui ressemble à : "
         "311791709280-xxxxx.apps.googleusercontent.com. C'est cette valeur "
         "qu'on met dans notre code."),
    ]

    for i, (titre, desc) in enumerate(etapes_google, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Étape {i} : {titre}")
        run.bold = True
        doc.add_paragraph(desc)

    p = doc.add_paragraph()
    run = p.add_run("Important : ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC1, 0x44, 0x0E)
    p.add_run(
        "Le Client ID n'est pas un secret. Il peut être visible dans le code HTML. "
        "C'est normal et sécurisé car Google vérifie aussi l'origine de la requête."
    )

    # 2.3
    doc.add_heading("2.3. Comment ça marche", level=2)
    doc.add_paragraph("Le processus se déroule en 5 étapes :")

    etapes_fonctionnement = [
        "Le site affiche le bouton « Se connecter avec Google » (fourni par Google).",
        "L'utilisateur clique et s'authentifie dans la fenêtre Google.",
        "Google renvoie un token JWT (un texte crypté contenant l'email et le nom).",
        "Notre code JavaScript reçoit ce token et l'envoie à notre serveur Django.",
        "Le serveur vérifie le token et stocke l'email Google dans la base de données.",
    ]

    for i, etape in enumerate(etapes_fonctionnement, 1):
        doc.add_paragraph(f"{i}. {etape}")

    p = doc.add_paragraph()
    run = p.add_run("Qu'est-ce qu'un JWT ? ")
    run.bold = True
    p.add_run(
        "Un JWT (JSON Web Token) est un texte en trois parties séparées par des points. "
        "La partie du milieu contient les informations (email, nom) encodées en Base64. "
        "La dernière partie est une signature qui prouve que le token vient bien de Google "
        "et n'a pas été modifié."
    )

    # 2.4
    doc.add_heading("2.4. Intégration dans le Tennis Club", level=2)
    doc.add_paragraph(
        "Dans notre projet, Google Identity Services est utilisé dans le flux de "
        "paiement de la cotisation. Quand un membre n'est pas en ordre de cotisation :"
    )

    doc.add_paragraph(
        "Une modale apparaît avec un bouton « Se connecter avec Google ».",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Le membre se connecte → son email Google est enregistré dans son profil.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Le bouton « Payer maintenant » apparaît (Stripe).",
        style='List Bullet'
    )
    doc.add_paragraph(
        "L'email Google est utilisé comme email pour le reçu Stripe.",
        style='List Bullet'
    )

    p = doc.add_paragraph()
    run = p.add_run("Fichiers concernés : ")
    run.bold = True
    p.add_run(
        "settings.py (GOOGLE_CLIENT_ID), views.py (vue google_auth), "
        "app.js (callback onGoogleSignIn), templates HTML (bouton Google)."
    )

    # 2.5
    doc.add_heading("2.5. Exemple de code", level=2)
    doc.add_paragraph(
        "Voici le code JavaScript côté navigateur pour afficher le bouton Google "
        "et envoyer le résultat au serveur :"
    )

    code_google_js = '''<!-- Charger le SDK Google Identity Services -->
<script src="https://accounts.google.com/gsi/client" async defer></script>

<!-- Bouton Google (généré automatiquement par Google) -->
<div id="g_id_onload"
     data-client_id="VOTRE_CLIENT_ID.apps.googleusercontent.com"
     data-callback="onGoogleSignIn">
</div>
<div class="g_id_signin" data-type="standard"></div>

<script>
// Cette fonction est appelée quand l'utilisateur se connecte
function onGoogleSignIn(response) {
    // Décoder le token pour récupérer l'email
    var payload = JSON.parse(atob(response.credential.split('.')[1]));

    // Envoyer l'email au serveur
    fetch('/google-auth/', {
        method: 'POST',
        headers: {
            'Content-Type': 'application/json',
            'X-CSRFToken': csrfToken
        },
        body: JSON.stringify({ email: payload.email })
    });
}
</script>'''

    p = doc.add_paragraph()
    run = p.add_run(code_google_js)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_paragraph()
    doc.add_paragraph(
        "Côté serveur (Python/Django), la vue qui reçoit l'email :"
    )

    code_google_py = '''@login_required
@require_POST
def google_auth(request):
    """Reçoit l'email Google et le stocke dans le profil"""
    body = json.loads(request.body)
    google_email = body.get('email')

    membre = Membre.objects.get(user=request.user)
    membre.google_email = google_email
    membre.save()

    return JsonResponse({'success': True, 'email': google_email})'''

    p = doc.add_paragraph()
    run = p.add_run(code_google_py)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_paragraph(
        "Un exemple complet et autonome est disponible dans le fichier "
        "exemples/exemple_google_auth.py."
    )

    doc.add_page_break()


def ajouter_section_stripe(doc):
    """Section 3 : Stripe Checkout"""
    doc.add_heading("3. API Stripe Checkout", level=1)

    # 3.1
    doc.add_heading("3.1. Qu'est-ce que c'est ?", level=2)
    doc.add_paragraph(
        "Stripe est une plateforme de paiement en ligne utilisée par des millions "
        "d'entreprises dans le monde (Amazon, Google, Shopify...). Stripe Checkout "
        "est un service de Stripe qui fournit une page de paiement prête à l'emploi, "
        "sécurisée et conforme aux normes bancaires (PCI DSS)."
    )
    doc.add_paragraph(
        "Concrètement, au lieu de créer nous-mêmes un formulaire de carte bancaire "
        "(ce qui est complexe et risqué), on redirige l'utilisateur vers une page Stripe "
        "qui gère tout : saisie de la carte, validation, sécurité 3D Secure, etc."
    )

    # 3.2
    doc.add_heading("3.2. Comment obtenir les clés API", level=2)
    doc.add_paragraph("Pour utiliser Stripe, il faut :")

    etapes_stripe = [
        ("Créer un compte Stripe",
         "Aller sur https://dashboard.stripe.com et s'inscrire (gratuit). "
         "Aucune carte bancaire n'est nécessaire pour le mode test."),
        ("Récupérer les clés de test",
         "Dans le Dashboard → « Developers » → « API keys ». "
         "Il y a deux clés : la clé publique (pk_test_...) et la clé secrète (sk_test_...)."),
        ("Configurer le webhook (optionnel)",
         "Dans « Developers » → « Webhooks » → « Add endpoint ». "
         "Entrer l'URL de notre serveur (ex: https://monsite.com/stripe-webhook/). "
         "Stripe nous donne un secret de webhook (whsec_...)."),
    ]

    for i, (titre, desc) in enumerate(etapes_stripe, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Étape {i} : {titre}")
        run.bold = True
        doc.add_paragraph(desc)

    # Tableau des clés
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Clé", "Commence par", "Rôle"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["Clé publique", "pk_test_...", "Utilisée côté navigateur (visible)"],
        ["Clé secrète", "sk_test_...", "Utilisée côté serveur uniquement (CONFIDENTIELLE)"],
        ["Secret webhook", "whsec_...", "Vérifie les notifications de Stripe"],
    ]
    for i, row_data in enumerate(data, 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val

    p = doc.add_paragraph()
    run = p.add_run("\nAttention : ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC1, 0x44, 0x0E)
    p.add_run(
        "La clé secrète (sk_test_...) ne doit JAMAIS être partagée ou mise dans du code "
        "visible par le navigateur. Elle permet de créer des paiements !"
    )

    # 3.3
    doc.add_heading("3.3. Comment ça marche", level=2)
    doc.add_paragraph("Le processus de paiement se déroule en 6 étapes :")

    etapes_paiement = [
        "Le membre clique sur « Payer maintenant » sur notre site.",
        "Notre serveur Django envoie une requête à l'API Stripe pour créer une session de paiement "
        "(montant, description, email du client, URLs de retour).",
        "Stripe crée la session et renvoie une URL unique.",
        "Notre serveur redirige le membre vers cette URL Stripe.",
        "Le membre entre sa carte bancaire sur la page Stripe (sécurisée).",
        "Après paiement : Stripe redirige vers notre page de succès ET envoie un webhook "
        "à notre serveur pour confirmer.",
    ]

    for i, etape in enumerate(etapes_paiement, 1):
        doc.add_paragraph(f"{i}. {etape}")

    p = doc.add_paragraph()
    run = p.add_run("Mode test : ")
    run.bold = True
    p.add_run(
        "En mode test, on utilise la carte fictive 4242 4242 4242 4242 "
        "avec n'importe quelle date future et CVC 123. Aucun vrai argent n'est débité."
    )

    # 3.4
    doc.add_heading("3.4. Intégration dans le Tennis Club", level=2)
    doc.add_paragraph(
        "Dans notre projet, Stripe Checkout est utilisé pour le paiement de la "
        "cotisation annuelle (50€). Le flux est le suivant :"
    )

    doc.add_paragraph(
        "Le membre se connecte avec Google (étape précédente).",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Il clique sur « Payer maintenant (50€) ».",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Le serveur crée une session Stripe avec le montant et l'email.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Le membre est redirigé vers la page de paiement Stripe.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Après paiement réussi, le champ « en_ordre_cotisation » passe à True "
        "et la date d'échéance est mise à jour (+ 1 an).",
        style='List Bullet'
    )

    p = doc.add_paragraph()
    run = p.add_run("Fichiers concernés : ")
    run.bold = True
    p.add_run(
        "settings.py (STRIPE_PUBLIC_KEY, STRIPE_SECRET_KEY), views.py "
        "(creer_session_paiement, stripe_webhook, paiement_succes), "
        "templates HTML (bouton Payer)."
    )

    # 3.5
    doc.add_heading("3.5. Exemple de code", level=2)
    doc.add_paragraph(
        "Voici le code Python côté serveur pour créer une session de paiement :"
    )

    code_stripe = '''import stripe
stripe.api_key = settings.STRIPE_SECRET_KEY

def creer_session_paiement(request):
    membre = Membre.objects.get(user=request.user)

    session = stripe.checkout.Session.create(
        mode='payment',
        line_items=[{
            'price_data': {
                'currency': 'eur',
                'unit_amount': 5000,  # 50€ en centimes
                'product_data': {
                    'name': 'Cotisation Tennis Club',
                },
            },
            'quantity': 1,
        }],
        customer_email=membre.google_email or membre.user.email,
        success_url=request.build_absolute_uri('/paiement/succes/'),
        cancel_url=request.build_absolute_uri('/paiement/annule/'),
    )

    return redirect(session.url)'''

    p = doc.add_paragraph()
    run = p.add_run(code_stripe)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_paragraph()
    doc.add_paragraph(
        "Le webhook qui traite la confirmation de paiement :"
    )

    code_webhook = '''@csrf_exempt
def stripe_webhook(request):
    payload = request.body
    sig = request.META.get('HTTP_STRIPE_SIGNATURE')

    event = stripe.Webhook.construct_event(
        payload, sig, settings.STRIPE_WEBHOOK_SECRET
    )

    if event['type'] == 'checkout.session.completed':
        session = event['data']['object']
        email = session.get('customer_email')
        # Mettre à jour le membre dans la base de données
        membre = Membre.objects.get(google_email=email)
        membre.en_ordre_cotisation = True
        membre.save()

    return HttpResponse(status=200)'''

    p = doc.add_paragraph()
    run = p.add_run(code_webhook)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'

    doc.add_paragraph(
        "Un exemple complet et autonome est disponible dans le fichier "
        "exemples/exemple_stripe.py."
    )

    doc.add_page_break()


def ajouter_resume(doc):
    """Section 4 : Résumé comparatif"""
    doc.add_heading("4. Résumé comparatif", level=1)

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["", "Google Identity Services", "Stripe Checkout"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h

    data = [
        ["Rôle", "Authentification / Lier un email", "Paiement en ligne"],
        ["Gratuit ?", "Oui (totalement)", "Oui (frais de 1.4% + 0.25€ par transaction)"],
        ["Clé nécessaire", "Client ID (publique)", "Clé publique + clé secrète"],
        ["Côté client", "SDK JavaScript de Google", "Redirection vers page Stripe"],
        ["Côté serveur", "Vérification du token JWT", "Création de session + webhook"],
        ["Données sensibles", "Aucune (le token est public)", "Clé secrète à protéger"],
        ["Carte de test", "N/A", "4242 4242 4242 4242"],
    ]

    for i, row_data in enumerate(data, 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val

    doc.add_paragraph()
    doc.add_paragraph(
        "Les deux APIs sont complémentaires : Google Identity Services identifie "
        "l'utilisateur (qui est-il ?), et Stripe Checkout traite le paiement "
        "(combien doit-il payer ?). Ensemble, elles permettent un flux de paiement "
        "complet et sécurisé pour la cotisation du Tennis Club."
    )


def main():
    doc = Document()

    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Générer le document
    ajouter_titre_page(doc)
    ajouter_table_des_matieres(doc)
    ajouter_introduction(doc)
    ajouter_section_google(doc)
    ajouter_section_stripe(doc)
    ajouter_resume(doc)

    # Sauvegarder
    output_dir = os.path.join(os.path.dirname(__file__), 'exemples')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'Documentation_APIs.docx')
    doc.save(output_path)
    print(f"Document généré : {output_path}")


if __name__ == '__main__':
    main()
